import os
import pypdf
import docx

def extract_text_from_file(file_path):
    """
    Extracts text from a given file path based on its extension.
    Supports: .txt, .docx, .pdf
    """
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.txt':
        return extract_txt(file_path)
    elif ext == '.docx':
        return extract_docx(file_path)
    elif ext == '.pdf':
        return extract_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. We support .txt, .docx, and .pdf files.")

def extract_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()

def extract_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

def extract_pdf(file_path):
    reader = pypdf.PdfReader(file_path)
    full_text = []
    
    for page_idx, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            full_text.append(page_text)
            
    return "\n".join(full_text)
